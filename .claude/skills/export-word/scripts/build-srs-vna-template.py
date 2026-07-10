# -*- coding: utf-8 -*-
"""
build-srs-vna-template.py — Dựng template `.claude/templates/word-reference-srs-vna.docx`
tu file SRS human-author "VNA.TOSS_SRS_System Admin_V0.1.docx" (bieu mau Viettel
"TAI LIEU THIET KE CHI TIET" — dong VNA.FIMS, Google Docs export).

Nguyen tac: styles.xml / theme1.xml lay NGUYEN VAN tu file nguon (fidelity by construction),
chi va 2 diem ky thuat:
  1. docDefaults <w:i w:val="1"/> -> "0"  (quirk Google Docs: default italic, moi run nguon
     deu tu set i=0; pandoc KHONG set i=0 tren run sinh ra -> phai sua default de render dung).
  2. Them style ImageCaption / TableCaption (pandoc dung; can giua, khong nghieng — theo
     quy uoc caption "Hinh N. / Bang N -" cua file nguon).

document.xml cua template = skeleton: trang bia {{placeholder}} + BANG GHI NHAN THAY DOI
+ TRANG KY + MUC LUC (field TOC) + mau moi style (Heading1-4 danh so, body, bang, caption).
sectPr cuoi = landscape KHONG header/footer refs (exporter tu inject de deterministic).

Chay:  python .claude/skills/export-word/scripts/build-srs-vna-template.py
Xem:   .claude/knowledge/srs-word-format-vna-tit.md (dac ta format day du)
"""
import zipfile, re, os, sys, io, shutil

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..', '..', '..'))
SRC = os.path.join(ROOT, 'ba', 'workspace', 'drafts', 'srs', 'VNA.TOSS_SRS_System Admin_V0.1.docx')
OUT = os.path.join(ROOT, '.claude', 'templates', 'word-reference-srs-vna.docx')
ASSETS = os.path.join(ROOT, '.claude', 'skills', 'export-word', 'assets', 'srs-vna')

W_NS = ('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture" '
        'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:w10="urn:schemas-microsoft-com:office:word"')

XML_DECL = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'

# ---------------------------------------------------------------- helpers
def r_bold(text):
    return f'<w:r><w:rPr><w:b w:val="1"/><w:bCs w:val="1"/></w:rPr><w:t xml:space="preserve">{text}</w:t></w:r>'

def r_plain(text):
    return f'<w:r><w:t xml:space="preserve">{text}</w:t></w:r>'

def p_center_bold(text, small_caps=False, before=120):
    sc = '<w:smallCaps w:val="1"/>' if small_caps else ''
    return (f'<w:p><w:pPr><w:spacing w:after="0" w:before="{before}" w:line="288" w:lineRule="auto"/>'
            f'<w:ind w:left="0" w:firstLine="0"/><w:jc w:val="center"/></w:pPr>'
            f'<w:r><w:rPr><w:b w:val="1"/><w:bCs w:val="1"/>{sc}</w:rPr>'
            f'<w:t xml:space="preserve">{text}</w:t></w:r></w:p>')

def p_blank(center=True):
    jc = '<w:jc w:val="center"/>' if center else ''
    return (f'<w:p><w:pPr><w:spacing w:after="0" w:before="120" w:line="288" w:lineRule="auto"/>'
            f'<w:ind w:left="0" w:firstLine="0"/>{jc}</w:pPr></w:p>')

def p_left(runs, first_line=270):
    return (f'<w:p><w:pPr><w:spacing w:after="0" w:before="120" w:line="288" w:lineRule="auto"/>'
            f'<w:ind w:left="0" w:firstLine="{first_line}"/><w:jc w:val="left"/></w:pPr>{runs}</w:p>')

def p_pagebreak():
    return '<w:p><w:pPr><w:spacing w:after="0" w:before="0"/></w:pPr><w:r><w:br w:type="page"/></w:r></w:p>'

def tc(text, width, shd=None, bold=False):
    shd_xml = f'<w:shd w:fill="{shd}" w:val="clear"/>' if shd else ''
    run = r_bold(text) if bold else r_plain(text)
    if not text:
        run = ''
    return (f'<w:tc><w:tcPr><w:tcW w:w="{width}" w:type="dxa"/>{shd_xml}'
            f'<w:tcMar><w:top w:w="60" w:type="dxa"/><w:left w:w="80" w:type="dxa"/>'
            f'<w:bottom w:w="60" w:type="dxa"/><w:right w:w="80" w:type="dxa"/></w:tcMar></w:tcPr>'
            f'<w:p><w:pPr><w:spacing w:after="0" w:before="40" w:line="240" w:lineRule="auto"/>'
            f'<w:ind w:left="0" w:firstLine="0"/><w:jc w:val="left"/></w:pPr>{run}</w:p></w:tc>')

# ---------------------------------------------------------------- front matter (voi placeholder)
CHANGE_GRID = [1080, 1080, 720, 1125, 1830, 2085, 1080]   # = 9000 twips (nguon: Data Maintenance)
CHANGE_HEAD = ['Ngày thay đổi', 'Vị trí thay đổi', 'A*, M, D', 'Nguồn gốc',
               'Phiên bản cũ', 'Mô tả thay đổi', 'Phiên bản mới']

def change_log_table(n_blank_rows=5):
    borders = ('<w:tblBorders>'
               + ''.join(f'<w:{side} w:color="000000" w:space="0" w:sz="4" w:val="dotted"/>'
                         for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'))
               + '</w:tblBorders>')
    grid = ''.join(f'<w:gridCol w:w="{w}"/>' for w in CHANGE_GRID)
    head = '<w:tr>' + ''.join(tc(h, w, shd='c0c0c0') for h, w in zip(CHANGE_HEAD, CHANGE_GRID)) + '</w:tr>'
    blank = '<w:tr>' + ''.join(tc('', w) for w in CHANGE_GRID) + '</w:tr>'
    return (f'<w:tbl><w:tblPr><w:tblW w:w="9000" w:type="dxa"/><w:jc w:val="left"/>'
            f'<w:tblInd w:w="80" w:type="dxa"/>{borders}<w:tblLayout w:type="fixed"/>'
            f'<w:tblLook w:val="0400"/></w:tblPr><w:tblGrid>{grid}</w:tblGrid>'
            + head + blank * n_blank_rows + '</w:tbl>')

def sign_block(label):
    tabs = '<w:r><w:tab/><w:tab/><w:tab/><w:tab/></w:r>'
    return (p_left(r_plain(label) + tabs + r_plain('&lt;Ngày&gt;'))
            + p_left(r_plain('&lt;Chức danh&gt;'))
            + p_blank(center=False) + p_blank(center=False) + p_blank(center=False))

def toc_field():
    instr = ' TOC \\h \\u \\z \\t &quot;Heading 1,1,Heading 2,2,Heading 3,3,&quot; '
    return ('<w:p><w:pPr><w:spacing w:after="0" w:before="120"/><w:ind w:left="0" w:firstLine="0"/>'
            '<w:jc w:val="left"/></w:pPr>'
            '<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
            f'<w:r><w:instrText xml:space="preserve">{instr}</w:instrText></w:r>'
            '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
            '<w:r><w:t xml:space="preserve">(Chọn toàn bộ tài liệu và nhấn F9 để cập nhật mục lục)</w:t></w:r>'
            '<w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>')

def cover_logo_para(rid='{{RID_LOGO}}'):
    return ('<w:p><w:pPr><w:spacing w:after="0" w:before="120" w:line="288" w:lineRule="auto"/>'
            '<w:ind w:left="0" w:right="0" w:firstLine="0"/><w:jc w:val="center"/></w:pPr>'
            '<w:r><w:drawing><wp:inline distB="0" distT="0" distL="114300" distR="114300">'
            '<wp:extent cx="2522220" cy="535940"/><wp:effectExtent b="0" l="0" r="0" t="0"/>'
            '<wp:docPr id="220" name="logo-viettel"/><a:graphic>'
            '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
            '<pic:pic><pic:nvPicPr><pic:cNvPr id="220" name="logo-viettel"/><pic:cNvPicPr/></pic:nvPicPr>'
            f'<pic:blipFill><a:blip r:embed="{rid}"/><a:stretch><a:fillRect/></a:stretch></pic:blipFill>'
            '<pic:spPr><a:xfrm><a:off x="0" y="0"/><a:ext cx="2522220" cy="535940"/></a:xfrm>'
            '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom></pic:spPr></pic:pic>'
            '</a:graphicData></a:graphic></wp:inline></w:drawing></w:r></w:p>')

def frontmatter(rid_logo='{{RID_LOGO}}'):
    """Fragment front matter — trang bia + bang ghi nhan thay doi + trang ky + muc luc."""
    parts = []
    parts.append(cover_logo_para(rid_logo))
    parts.append(p_center_bold('TẬP ĐOÀN CÔNG NGHIỆP - VIỄN THÔNG QUÂN ĐỘI VIETTEL'))
    parts.append(p_center_bold('{{DON_VI}}'))
    parts += [p_blank()] * 4
    parts.append(p_center_bold('BIỂU MẪU'))
    parts.append(p_center_bold('TÀI LIỆU THIẾT KẾ CHI TIẾT'))
    parts += [p_blank()] * 2
    parts.append(p_left(r_plain('Mã hiệu dự án: {{MA_HIEU_DU_AN}}')))
    parts.append(p_left(r_plain('Mã hiệu tài liệu: {{MA_HIEU_TAI_LIEU}}')))
    parts += [p_blank(center=False)] * 9
    parts.append(p_left(r_plain('          '
                                '          '
                                '&lt;Hà Nội, {{THANG_NAM}}&gt;')))
    parts.append(p_pagebreak())
    parts.append(p_center_bold('BẢNG GHI NHẬN THAY ĐỔI TÀI LIỆU', small_caps=True, before=240))
    parts.append(p_blank(center=False))
    parts.append(change_log_table())
    parts.append(p_left(r_plain('*A – Tạo mới, M – Sửa đổi, D – Xóa bỏ'), first_line=0))
    parts.append(p_pagebreak())
    parts.append(p_center_bold('TRANG KÝ', small_caps=True, before=240))
    parts.append(p_blank(center=False))
    parts.append(sign_block('Người lập:'))
    parts.append(sign_block('Người xem xét:'))
    parts.append(sign_block('Người xem xét:'))
    parts.append(sign_block('Người phê duyệt:'))
    parts.append(p_pagebreak())
    parts.append(p_center_bold('MỤC LỤC', small_caps=True, before=240))
    parts.append(toc_field())
    return ''.join(parts)

# ---------------------------------------------------------------- sectPr (2 khô? trang)
SECTPR_PORTRAIT = (  # front matter + phan A (refs do exporter dien {{...}})
    '<w:sectPr>{{HDRFTR_REFS}}'
    '<w:pgSz w:h="16838" w:w="11906" w:orient="portrait"/>'
    '<w:pgMar w:bottom="1440" w:top="1440" w:left="1797" w:right="1440" w:header="567" w:footer="567"/>'
    '<w:pgNumType w:start="1"/><w:titlePg w:val="1"/></w:sectPr>')

SECTPR_LANDSCAPE = (  # phan than tai lieu (khong refs — ke thua section truoc)
    '<w:sectPr><w:type w:val="nextPage"/>'
    '<w:pgSz w:h="11906" w:w="16838" w:orient="landscape"/>'
    '<w:pgMar w:bottom="1440" w:top="1797" w:left="1276" w:right="1440" w:header="567" w:footer="567"/>'
    '</w:sectPr>')

# ---------------------------------------------------------------- headers/footers (lam sach tu nguon)
def _hf(kind, body):
    tag = 'hdr' if kind == 'h' else 'ftr'
    return f'{XML_DECL}<w:{tag} {W_NS}>{body}</w:{tag}>'

HEADER_FIRST = _hf('h', '<w:p><w:pPr><w:ind w:firstLine="270"/></w:pPr></w:p>')  # trang bia: trong
HEADER_EVEN = _hf('h',  # trang chan: dong ke duoi (evenAndOddHeaders off -> khong render, giu theo nguon)
    '<w:p><w:pPr><w:pBdr><w:bottom w:color="000000" w:space="1" w:sz="4" w:val="single"/></w:pBdr>'
    '<w:tabs><w:tab w:val="right" w:leader="none" w:pos="9000"/></w:tabs>'
    '<w:spacing w:after="0" w:before="120" w:line="240" w:lineRule="auto"/>'
    '<w:ind w:left="0" w:right="14" w:firstLine="0"/><w:jc w:val="both"/>'
    '<w:rPr><w:color w:val="000000"/><w:sz w:val="22"/><w:szCs w:val="22"/></w:rPr></w:pPr></w:p>')
_FT_RPR = ('<w:rPr><w:rFonts w:ascii="Times New Roman" w:cs="Times New Roman" '
           'w:eastAsia="Times New Roman" w:hAnsi="Times New Roman"/>'
           '<w:i w:val="1"/><w:iCs w:val="1"/><w:color w:val="000000"/>'
           '<w:sz w:val="22"/><w:szCs w:val="22"/></w:rPr>')
_FT_PPR = ('<w:pPr><w:tabs><w:tab w:val="center" w:leader="none" w:pos="4320"/>'
           '<w:tab w:val="right" w:leader="none" w:pos="9000"/></w:tabs>'
           '<w:spacing w:after="0" w:before="120" w:line="360" w:lineRule="auto"/>'
           '<w:ind w:left="0" w:firstLine="0"/><w:jc w:val="left"/>'
           '<w:rPr><w:color w:val="000000"/><w:sz w:val="22"/><w:szCs w:val="22"/></w:rPr></w:pPr>')
FOOTER_DEFAULT = _hf('f',  # "trang/tong" nghieng TNR 11pt — theo nguon System Admin
    f'<w:p>{_FT_PPR}'
    f'<w:r>{_FT_RPR}<w:fldChar w:fldCharType="begin"/></w:r>'
    f'<w:r>{_FT_RPR}<w:instrText xml:space="preserve">PAGE</w:instrText></w:r>'
    f'<w:r>{_FT_RPR}<w:fldChar w:fldCharType="end"/></w:r>'
    f'<w:r>{_FT_RPR}<w:t>/</w:t></w:r>'
    f'<w:r>{_FT_RPR}<w:fldChar w:fldCharType="begin"/></w:r>'
    f'<w:r>{_FT_RPR}<w:instrText xml:space="preserve">NUMPAGES</w:instrText></w:r>'
    f'<w:r>{_FT_RPR}<w:fldChar w:fldCharType="end"/></w:r></w:p>')
FOOTER_FIRST = _hf('f', f'<w:p>{_FT_PPR}</w:p>')  # trang bia: khong so trang
FOOTER_EVEN = _hf('f',
    f'<w:p>{_FT_PPR}'
    f'<w:r>{_FT_RPR}<w:fldChar w:fldCharType="begin"/></w:r>'
    f'<w:r>{_FT_RPR}<w:instrText xml:space="preserve">PAGE</w:instrText></w:r>'
    f'<w:r>{_FT_RPR}<w:fldChar w:fldCharType="end"/></w:r></w:p>')

HDRFTR_PARTS = {  # ten part -> (loai ref, noi dung)
    'header1.xml': ('headerReference', 'first', HEADER_FIRST),
    'header2.xml': ('headerReference', 'even', HEADER_EVEN),
    'footer1.xml': ('footerReference', 'default', FOOTER_DEFAULT),
    'footer2.xml': ('footerReference', 'first', FOOTER_FIRST),
    'footer3.xml': ('footerReference', 'even', FOOTER_EVEN),
}

# ---------------------------------------------------------------- numbering (heading multilevel)
def heading_abstract_num(aid):
    """abstractNum multilevel danh so heading — nguyen van abs 15 cua file nguon (lam tron twips)."""
    lvls = []
    ind = [(360, 360), (576, 576), (720, 720), (850, 870), (1008, 1008),
           (1152, 1152), (1296, 1296), (1440, 1440), (1584, 1584)]
    for i in range(9):
        lvltext = '.'.join(f'%{j+1}' for j in range(i + 1)) + ('.' if i == 0 else '')
        rpr = '<w:rPr/>'
        if i == 2:  # nguon: so cua Heading3 nghieng, 12pt
            rpr = '<w:rPr><w:i w:val="1"/><w:iCs w:val="1"/><w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr>'
        left, hang = ind[i]
        lvls.append(f'<w:lvl w:ilvl="{i}"><w:start w:val="1"/><w:numFmt w:val="decimal"/>'
                    f'<w:lvlText w:val="{lvltext}"/><w:lvlJc w:val="left"/>'
                    f'<w:pPr><w:ind w:left="{left}" w:hanging="{hang}"/></w:pPr>{rpr}</w:lvl>')
    return f'<w:abstractNum w:abstractNumId="{aid}">' + ''.join(lvls) + '</w:abstractNum>'

TEMPLATE_NUMBERING = (f'{XML_DECL}<w:numbering {W_NS}>'
                      + heading_abstract_num(15)
                      + '<w:num w:numId="15"><w:abstractNumId w:val="15"/></w:num></w:numbering>')

# ---------------------------------------------------------------- skeleton samples
def sample_heading(level, text, ilvl):
    return (f'<w:p><w:pPr><w:pStyle w:val="Heading{level}"/>'
            f'<w:numPr><w:ilvl w:val="{ilvl}"/><w:numId w:val="15"/></w:numPr></w:pPr>'
            f'{r_plain(text)}</w:p>')

def sample_table():
    borders = ('<w:tblBorders>'
               + ''.join(f'<w:{s} w:color="000000" w:space="0" w:sz="4" w:val="single"/>'
                         for s in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'))
               + '</w:tblBorders>')
    g = [2000, 4000, 6000]
    grid = ''.join(f'<w:gridCol w:w="{w}"/>' for w in g)
    head = '<w:tr>' + ''.join(tc(h, w, shd='f2f2f2', bold=True) for h, w in zip(['STT', 'Trường', 'Mô tả'], g)) + '</w:tr>'
    row = '<w:tr>' + ''.join(tc(t, w) for t, w in zip(['1', '(mẫu)', '(mẫu bảng mô tả — viền đơn đen 0.5pt, header xám f2f2f2)'], g)) + '</w:tr>'
    return (f'<w:tbl><w:tblPr><w:tblW w:w="12000" w:type="dxa"/><w:jc w:val="left"/>{borders}'
            f'<w:tblLayout w:type="fixed"/><w:tblLook w:val="0400"/></w:tblPr>'
            f'<w:tblGrid>{grid}</w:tblGrid>{head}{row}</w:tbl>')

def skeleton_body():
    b = []
    b.append(frontmatter(rid_logo='rIdLogo'))
    # phan A (portrait) — heading part khong danh so
    b.append('<w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr>' + r_plain('A - THÔNG TIN CHUNG') + '</w:p>')
    b.append(sample_heading(1, 'GIỚI THIỆU', 0))
    b.append(sample_heading(2, 'Mục đích', 1))
    b.append('<w:p><w:pPr><w:ind w:left="270" w:firstLine="0"/></w:pPr>'
             + r_plain('(Mẫu đoạn thân bài — kế thừa docDefaults: 12pt, giãn dòng 1.5, căn đều 2 lề, thụt trái 270 twips.)')
             + '</w:p>')
    b.append(sample_heading(3, 'Mẫu Heading 3', 2))
    b.append(sample_heading(4, 'Mẫu Heading 4', 3))
    # section break -> landscape
    b.append('<w:p><w:pPr>' + SECTPR_PORTRAIT.replace('{{HDRFTR_REFS}}', '{{TPL_REFS}}') + '</w:pPr></w:p>')
    b.append('<w:p><w:pPr><w:pStyle w:val="Heading1"/></w:pPr>' + r_plain('B - THIẾT KẾ CHI TIẾT') + '</w:p>')
    b.append(sample_table())
    b.append('<w:p><w:pPr><w:pStyle w:val="ImageCaption"/></w:pPr>' + r_plain('Hình 1. (Mẫu caption hình — căn giữa)') + '</w:p>')
    return ''.join(b)

# ---------------------------------------------------------------- build
def main():
    os.makedirs(ASSETS, exist_ok=True)
    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    src = zipfile.ZipFile(SRC)

    styles = src.read('word/styles.xml').decode('utf-8')
    # (1) quirk gdocs: default italic -> tat (run nguon nao cung tu set i=0; pandoc thi khong)
    styles = styles.replace('<w:rPrDefault><w:rPr><w:i w:val="1"/><w:iCs w:val="1"/>',
                            '<w:rPrDefault><w:rPr><w:i w:val="0"/><w:iCs w:val="0"/>', 1)
    # (2) them style caption cho pandoc (can giua, khong nghieng — theo quy uoc nguon)
    cap = ('<w:style w:type="paragraph" w:styleId="ImageCaption"><w:name w:val="Image Caption"/>'
           '<w:basedOn w:val="Normal"/><w:pPr><w:spacing w:after="120" w:before="60"/>'
           '<w:ind w:left="0" w:firstLine="0"/><w:jc w:val="center"/></w:pPr>'
           '<w:rPr><w:i w:val="0"/><w:iCs w:val="0"/></w:rPr></w:style>'
           '<w:style w:type="paragraph" w:styleId="TableCaption"><w:name w:val="Table Caption"/>'
           '<w:basedOn w:val="Normal"/><w:pPr><w:spacing w:after="60" w:before="120"/>'
           '<w:ind w:left="0" w:firstLine="0"/><w:jc w:val="left"/></w:pPr>'
           '<w:rPr><w:i w:val="0"/><w:iCs w:val="0"/></w:rPr></w:style>')
    if 'ImageCaption' not in styles:
        styles = styles.replace('</w:styles>', cap + '</w:styles>')

    theme = src.read('word/theme/theme1.xml').decode('utf-8')
    settings = src.read('word/settings.xml').decode('utf-8')

    # logo trang bia: drawing dau tien cua document -> media tuong ung
    doc_src = src.read('word/document.xml').decode('utf-8')
    rels_src = src.read('word/_rels/document.xml.rels').decode('utf-8')
    emb = re.search(r'<w:drawing>.*?r:embed="(rId\d+)"', doc_src, re.S).group(1)
    tgt = re.search(rf'Id="{emb}"[^>]*Target="([^"]+)"', rels_src).group(1)
    logo = src.read('word/' + tgt)
    src.close()

    with open(os.path.join(ASSETS, 'logo-viettel-cover.png'), 'wb') as f:
        f.write(logo)
    with open(os.path.join(ASSETS, 'frontmatter.xml'), 'w', encoding='utf-8') as f:
        f.write(frontmatter())          # fragment voi {{PLACEHOLDER}} cho exporter
    with open(os.path.join(ASSETS, 'sectpr-portrait.xml'), 'w', encoding='utf-8') as f:
        f.write(SECTPR_PORTRAIT)
    with open(os.path.join(ASSETS, 'sectpr-landscape.xml'), 'w', encoding='utf-8') as f:
        f.write(SECTPR_LANDSCAPE)
    for name, (_, _, content) in HDRFTR_PARTS.items():
        with open(os.path.join(ASSETS, name), 'w', encoding='utf-8') as f:
            f.write(content)
    with open(os.path.join(ASSETS, 'numbering-heading.xml'), 'w', encoding='utf-8') as f:
        f.write(heading_abstract_num(900).replace('"900"', '"{{AID}}"'))

    # ------- template docx -------
    tpl_refs = ''  # trong template skeleton: khong refs (exporter inject sau)
    document = (f'{XML_DECL}<w:document {W_NS}><w:body>'
                + skeleton_body().replace('{{TPL_REFS}}', tpl_refs)
                + SECTPR_LANDSCAPE + '</w:body></w:document>')

    content_types = (f'{XML_DECL}<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                     '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                     '<Default Extension="xml" ContentType="application/xml"/>'
                     '<Default Extension="png" ContentType="image/png"/>'
                     '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                     '<Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
                     '<Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>'
                     '<Override PartName="/word/settings.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.settings+xml"/>'
                     '<Override PartName="/word/theme/theme1.xml" ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/>'
                     '<Override PartName="/docProps/core.xml" ContentType="application/vnd.openxmlformats-package.core-properties+xml"/>'
                     '<Override PartName="/docProps/app.xml" ContentType="application/vnd.openxmlformats-officedocument.extended-properties+xml"/>'
                     '</Types>')
    root_rels = (f'{XML_DECL}<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                 '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
                 '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties" Target="docProps/core.xml"/>'
                 '<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties" Target="docProps/app.xml"/>'
                 '</Relationships>')
    doc_rels = (f'{XML_DECL}<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
                '<Relationship Id="rId2" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering" Target="numbering.xml"/>'
                '<Relationship Id="rId3" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/settings" Target="settings.xml"/>'
                '<Relationship Id="rId4" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme" Target="theme/theme1.xml"/>'
                '<Relationship Id="rIdLogo" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image" Target="media/logo.png"/>'
                '</Relationships>')
    core = (f'{XML_DECL}<cp:coreProperties '
            'xmlns:cp="http://schemas.openxmlformats.org/package/2006/metadata/core-properties" '
            'xmlns:dc="http://purl.org/dc/elements/1.1/" xmlns:dcterms="http://purl.org/dc/terms/" '
            'xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance">'
            '<dc:title>word-reference-srs-vna</dc:title>'
            '<dc:description>Template Word SRS theo bieu mau Viettel TAI LIEU THIET KE CHI TIET (dong VNA.FIMS/TOSS)</dc:description>'
            '</cp:coreProperties>')
    app = (f'{XML_DECL}<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/extended-properties" '
           'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
           '<Application>build-srs-vna-template.py</Application></Properties>')

    if os.path.exists(OUT):
        os.remove(OUT)
    zout = zipfile.ZipFile(OUT, 'w', zipfile.ZIP_DEFLATED)
    # python zipfile ghi entry voi '/' (OPC hop le) — KHONG dung .NET CreateFromDirectory
    zout.writestr('[Content_Types].xml', content_types)
    zout.writestr('_rels/.rels', root_rels)
    zout.writestr('word/document.xml', document)
    zout.writestr('word/_rels/document.xml.rels', doc_rels)
    zout.writestr('word/styles.xml', styles)
    zout.writestr('word/numbering.xml', TEMPLATE_NUMBERING)
    zout.writestr('word/settings.xml', settings)
    zout.writestr('word/theme/theme1.xml', theme)
    zout.writestr('word/media/logo.png', logo)
    zout.writestr('docProps/core.xml', core)
    zout.writestr('docProps/app.xml', app)
    zout.close()

    # smoke checks
    import xml.dom.minidom
    z = zipfile.ZipFile(OUT)
    for n in z.namelist():
        if n.endswith('.xml') or n.endswith('.rels'):
            xml.dom.minidom.parseString(z.read(n))
    assert not any('\\' in n for n in z.namelist()), 'entry backslash!'
    z.close()
    from docx import Document
    d = Document(OUT)
    print('TEMPLATE OK:', OUT)
    print('  paragraphs in skeleton:', len(d.paragraphs), '| tables:', len(d.tables))
    print('  assets ->', ASSETS)

if __name__ == '__main__':
    main()
